import docx

def main():
    doc = docx.Document("Project Documentation.docx")
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    
    # Also get text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            fullText.append(" | ".join(row_text))

    with open("Project_Documentation.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(fullText))
    print("Extracted successfully!")

if __name__ == "__main__":
    main()
